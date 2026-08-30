#!/usr/bin/env python3
"""Generate the sample CV fixture used by the parser tests.

The real `Mohammed_Alostah_CV.docx` is not in the repository — a CV is personal
data and does not belong in version control. This builds an equivalent synthetic
document from the profile seed in `job-agent-plan.md` section 2, so the tests
exercise the same shapes (a skills block, tabular experience, dated roles)
without committing anyone's actual history.

To work with the real CV, drop it in `fixtures/resumes/` locally; it is
gitignored.
"""

from __future__ import annotations

import pathlib

from docx import Document

OUTPUT = (
    pathlib.Path(__file__).resolve().parent.parent
    / "fixtures"
    / "resumes"
    / "sample_engineering_lead.docx"
)

SKILLS = [
    "Python",
    "FastAPI",
    "Django",
    "Node.js",
    "Express.js",
    "PHP",
    "Laravel",
    "PostgreSQL",
    "MySQL",
    "MongoDB",
    "Redis",
    "Elasticsearch",
    "Azure",
    "AKS",
    "AWS",
    "Kubernetes",
    "Docker",
    "Terraform",
    "OAuth 2.0",
    "Azure AD SSO",
    "JWT",
    "STRIDE",
    "LangChain",
    "LangSmith",
    "CrewAI",
    "RAG",
]

ROLES = [
    (
        "Engineering Lead",
        "Northwind Systems",
        "Amman, Jordan",
        "Feb 2022 - Present",
        [
            "Led a team of six engineers delivering a multi-tenant SaaS platform.",
            "Cut p95 API latency from 840 ms to 210 ms by reworking the query layer.",
            "Introduced architecture reviews that reduced production incidents by 40 percent.",
        ],
    ),
    (
        "Senior Backend Engineer",
        "Cedar Analytics",
        "Amman, Jordan",
        "Jun 2019 - Jan 2022",
        [
            "Designed an event-driven ingestion pipeline handling 12 million records per day.",
            "Migrated a monolith to microservices on AKS with zero planned downtime.",
        ],
    ),
    (
        "Backend Engineer",
        "Levant Web Works",
        "Amman, Jordan",
        "Sep 2017 - May 2019",
        [
            "Built REST APIs in Laravel and Express.js for regional e-commerce clients.",
        ],
    ),
]


def main() -> None:
    document = Document()

    document.add_paragraph("Sample Candidate")
    document.add_paragraph("Engineering Lead / Senior Backend Engineer")
    document.add_paragraph("Amman, Jordan | sample.candidate@example.com | +962 79 000 0000")
    document.add_paragraph("github.com/sample-candidate")

    document.add_paragraph("Professional Summary")
    document.add_paragraph(
        "Engineering lead with 7 years of experience building distributed systems, "
        "microservices, and multi-tenant SaaS platforms. Arabic native, English fluent."
    )

    document.add_paragraph("Technical Skills")
    document.add_paragraph(", ".join(SKILLS))

    document.add_paragraph("Experience")
    # Deliberately tabular: many CVs lay experience out this way, and a parser
    # that only reads paragraphs loses all of it.
    table = document.add_table(rows=0, cols=2)
    for title, company, location, dates, achievements in ROLES:
        row = table.add_row().cells
        row[0].text = f"{title}, {company}"
        row[1].text = f"{location} | {dates}"
        bullets = table.add_row().cells
        bullets[0].text = ""
        bullets[1].text = "\n".join(achievements)

    document.add_paragraph("Education")
    document.add_paragraph("BSc Computer Science, University of Jordan, 2017")

    document.add_paragraph("Certifications")
    document.add_paragraph("Certified Kubernetes Application Developer")
    document.add_paragraph("Microsoft Certified: Azure Solutions Architect Expert")

    document.add_paragraph("Languages")
    document.add_paragraph("Arabic (native), English (fluent)")

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    document.save(OUTPUT)
    print(f"wrote {OUTPUT}")


if __name__ == "__main__":
    main()
