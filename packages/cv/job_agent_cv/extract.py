"""Deterministic text extraction from uploaded CVs.

Deterministic on purpose: the same file always yields the same text, no model
involved. Everything downstream — fact evidence, reprocessing, golden tests —
depends on that being true.
"""

from __future__ import annotations

import io
import re
from collections.abc import Iterator
from dataclasses import dataclass, field
from enum import StrEnum

from docx import Document
from docx.document import Document as DocxDocument
from docx.oxml.ns import qn
from docx.table import Table
from docx.text.paragraph import Paragraph
from pypdf import PdfReader
from pypdf.errors import PdfReadError

from job_agent_cv.errors import (
    EmptyDocumentError,
    ExtractionError,
    FileTooLargeError,
    UnsupportedFormatError,
)

#: 10 MB. A CV that exceeds this is either not a CV or is full of images.
MAX_UPLOAD_BYTES = 10 * 1024 * 1024

#: Below this, treat the document as empty: a scanned PDF with no text layer
#: parses fine and yields a handful of stray characters.
MIN_USEFUL_CHARS = 120

DOCX_MAGIC = b"PK\x03\x04"
PDF_MAGIC = b"%PDF-"


class ResumeFormat(StrEnum):
    DOCX = "docx"
    PDF = "pdf"


#: Headings we recognise, mapped to a canonical section name. Matching is on the
#: line alone, so a mention of "education" mid-sentence does not split a CV.
_SECTION_PATTERNS: tuple[tuple[str, re.Pattern[str]], ...] = (
    ("summary", re.compile(r"^(professional\s+)?(summary|profile|about\s+me|objective)\b", re.I)),
    (
        "experience",
        re.compile(r"^(work\s+|professional\s+|employment\s+)?(experience|history)\b", re.I),
    ),
    ("education", re.compile(r"^education(\s+and\s+training)?\b", re.I)),
    (
        "skills",
        re.compile(
            r"^(technical\s+|core\s+|key\s+)?(skills|competencies|technologies|tech\s+stack)\b",
            re.I,
        ),
    ),
    ("certifications", re.compile(r"^(certifications?|licenses?|credentials)\b", re.I)),
    ("projects", re.compile(r"^(projects?|selected\s+work|portfolio)\b", re.I)),
    ("languages", re.compile(r"^languages?\b", re.I)),
    (
        "contact",
        re.compile(r"^(contact|contact\s+details|personal\s+details)\b", re.I),
    ),
)

#: Includes U+00A0: DOCX exports are full of non-breaking spaces.
_WHITESPACE_RE = re.compile("[ \t\u00a0]+")
_BLANK_LINES_RE = re.compile(r"\n{3,}")


@dataclass(frozen=True, slots=True)
class ExtractedDocument:
    """Text plus the structure we could recover deterministically."""

    text: str
    format: ResumeFormat
    #: Canonical section name to its text. Missing sections are simply absent.
    sections: dict[str, str] = field(default_factory=dict)
    page_count: int | None = None

    @property
    def char_count(self) -> int:
        return len(self.text)


def sniff_format(data: bytes, *, filename: str | None = None) -> ResumeFormat:
    """Identify the format from content, not from the extension.

    A ``.docx`` that is really a renamed executable must not reach a parser, so
    the magic bytes decide and the filename is only a tie-breaker in the error.
    """
    if data.startswith(PDF_MAGIC):
        return ResumeFormat.PDF
    if data.startswith(DOCX_MAGIC):
        # A DOCX is a zip; so is an XLSX. The parser rejects the wrong contents.
        return ResumeFormat.DOCX
    hint = f" (filename: {filename})" if filename else ""
    raise UnsupportedFormatError(f"upload is neither a PDF nor a DOCX{hint}")


def normalise(text: str) -> str:
    """Collapse whitespace without losing line structure."""
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    text = "".join(character for character in text if character.isprintable() or character == "\n")
    lines = [_WHITESPACE_RE.sub(" ", line).strip() for line in text.split("\n")]
    return _BLANK_LINES_RE.sub("\n\n", "\n".join(lines)).strip()


def _iter_blocks(document: DocxDocument) -> Iterator[Paragraph | Table]:
    """Yield paragraphs and tables in document order.

    ``document.paragraphs`` and ``document.tables`` are separate sequences, so
    reading one then the other puts every table at the end of the text. In a CV
    that lays experience out in a table, the roles then land after the education
    heading and the experience section reads as empty.
    """
    for child in document.element.body.iterchildren():
        if child.tag == qn("w:p"):
            yield Paragraph(child, document)
        elif child.tag == qn("w:tbl"):
            yield Table(child, document)


def _table_text(table: Table) -> list[str]:
    lines: list[str] = []
    for row in table.rows:
        cells = [cell.text.strip() for cell in row.cells]
        # A merged cell repeats its text across the span; collapse the repeats.
        deduped: list[str] = []
        for cell in cells:
            if cell and (not deduped or deduped[-1] != cell):
                deduped.append(cell)
        if deduped:
            lines.append(" | ".join(deduped))
    return lines


def _docx_text(data: bytes) -> str:
    try:
        document = Document(io.BytesIO(data))
    except Exception as exc:  # python-docx raises a variety of types
        raise ExtractionError(f"could not read DOCX: {exc}") from exc

    parts: list[str] = []
    for block in _iter_blocks(document):
        if isinstance(block, Paragraph):
            parts.append(block.text)
        else:
            parts.extend(_table_text(block))
    return "\n".join(parts)


def _pdf_text(data: bytes) -> tuple[str, int]:
    try:
        reader = PdfReader(io.BytesIO(data))
        pages = [page.extract_text() or "" for page in reader.pages]
    except (PdfReadError, ValueError, OSError) as exc:
        raise ExtractionError(f"could not read PDF: {exc}") from exc
    return "\n".join(pages), len(pages)


def split_sections(text: str) -> dict[str, str]:
    """Split on recognised headings. Text before the first heading is ignored
    here; the caller still has the full document."""
    sections: dict[str, list[str]] = {}
    current: str | None = None

    for line in text.split("\n"):
        stripped = line.strip()
        if stripped and len(stripped) <= 60:
            matched = next(
                (name for name, pattern in _SECTION_PATTERNS if pattern.match(stripped)), None
            )
            if matched is not None:
                current = matched
                sections.setdefault(current, [])
                continue
        if current is not None:
            sections[current].append(line)

    return {name: "\n".join(lines).strip() for name, lines in sections.items() if any(lines)}


def extract(
    data: bytes, *, filename: str | None = None, max_bytes: int = MAX_UPLOAD_BYTES
) -> ExtractedDocument:
    """Extract text from an uploaded CV.

    Raises rather than returning partial results: a half-read CV would produce
    a profile with silently missing experience.
    """
    if not data:
        raise EmptyDocumentError("uploaded file is empty")
    if len(data) > max_bytes:
        raise FileTooLargeError(f"upload is {len(data)} bytes; the limit is {max_bytes}")

    resume_format = sniff_format(data, filename=filename)
    page_count: int | None = None
    if resume_format is ResumeFormat.PDF:
        raw, page_count = _pdf_text(data)
    else:
        raw = _docx_text(data)

    text = normalise(raw)
    if len(text) < MIN_USEFUL_CHARS:
        raise EmptyDocumentError(
            "document contains no readable text; if it is a scan, it needs OCR before upload"
        )

    return ExtractedDocument(
        text=text,
        format=resume_format,
        sections=split_sections(text),
        page_count=page_count,
    )
